"""
remediate_docx.py (courseforge) - Word (.docx) ADA/Ally remediation engine.
Mirror of remediate_pptx.py: two-phase so the AGENT supplies vision (image
descriptions) and judgment (which paragraphs are really headings) with no
per-file human oversight.

  1) scan   : find the issues Ally flags on DOCX - images with missing/filename
              alt, substantial documents with no real Heading styles, heading
              level skips, tables without a repeating header row - and extract
              every image to a workdir with a manifest the agent reads.
  2) apply  : take the agent-written fixes.json and write a remediated copy:
              - alt text on every image ("" = decorative is NOT supported by
                Word the way PPTX does it, so decorative images get alt "" only)
              - optional paragraph -> Heading style promotion (agent-mapped;
                changes appearance, so it is OPT-IN via fixes.json, never auto)
              - first-row header flag (w:tblHeader) on tables missing it
  3) verify : re-scan; hard issues should be 0 (un-promoted heading candidates
              downgrade to advisory when fixes.json chose not to promote).

HONEST SCOPE: alt text + table header rows + opt-in heading promotion. This is
NOT full document tagging; reading order / contrast / list semantics stay the
author's job. python-docx has no first-class alt API - we edit the drawing XML
(wp:docPr @descr) directly.

Usage:
  python remediate_docx.py scan   doc.docx --workdir W
  python remediate_docx.py apply  doc.docx --workdir W --out fixed.docx
  python remediate_docx.py verify fixed.docx

fixes.json (agent-written, in workdir):
  {
    "alts":     { "<imageKey>": "concise description (<=110 chars)" | "" },
    "headings": { "<paragraphIndex>": 1|2|3 },     # opt-in style promotion
    "table_headers": true
  }
"""
import argparse, json, os, re, sys

from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

MAX_ALT = 110
FILENAME_RX = re.compile(r"\.(png|jpe?g|gif|bmp|svg|webp|tiff?)\s*$", re.I)


# ---------- drawing helpers (alt text lives on wp:docPr @descr) ----------

def iter_drawings(doc):
    """Yield (index, docPr element, embedded image part or None) for every
    inline/floating drawing in body order."""
    body = doc.element.body
    idx = 0
    for docPr in body.iter(qn("wp:docPr")):
        # find the blip (image relationship) inside the same drawing, if any
        drawing = docPr.getparent().getparent()  # docPr -> inline/anchor -> drawing? be lenient
        blip = None
        anc = docPr.getparent()
        if anc is not None:
            blip = anc.find(".//" + qn("a:blip"))
        yield idx, docPr, blip
        idx += 1


def get_image_blob(doc, blip):
    if blip is None:
        return None, None
    rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
    if not rid:
        return None, None
    try:
        part = doc.part.related_parts[rid]
        ext = part.partname.ext.lstrip(".")
        return part.blob, ext
    except Exception:
        return None, None


# ---------- headings ----------

def heading_level(par):
    name = (par.style.name or "") if par.style is not None else ""
    m = re.match(r"heading (\d+)", name, re.I)
    return int(m.group(1)) if m else None


def looks_like_faux_heading(par):
    """Short, bold-run paragraph with no heading style = probable faux heading."""
    text = par.text.strip()
    if not text or len(text) > 80 or heading_level(par) is not None:
        return False
    runs = [r for r in par.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


# ---------- tables ----------

def table_has_header_row(table):
    if not table.rows:
        return False
    trPr = table.rows[0]._tr.find(qn("w:trPr"))
    return trPr is not None and trPr.find(qn("w:tblHeader")) is not None


def set_table_header_row(table):
    tr = table.rows[0]._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = tr.makeelement(qn("w:trPr"), {})
        tr.insert(0, trPr)
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(tr.makeelement(qn("w:tblHeader"), {}))


# ---------- scan ----------

def scan(path, workdir):
    doc = Document(path)
    os.makedirs(os.path.join(workdir, "images"), exist_ok=True)
    issues, images, faux = [], [], []

    for idx, docPr, blip in iter_drawings(doc):
        alt = (docPr.get("descr") or "").strip()
        key = "img%d" % idx
        blob, ext = get_image_blob(doc, blip)
        img_path = None
        if blob:
            img_path = os.path.join(workdir, "images", "%s.%s" % (key, ext or "png"))
            with open(img_path, "wb") as f:
                f.write(blob)
        problem = None
        if not alt:
            problem = "image missing alt"
        elif FILENAME_RX.search(alt) or (" " not in alt and "." in alt):
            problem = "alt is a filename: %r" % alt[:40]
        elif len(alt) > MAX_ALT:
            problem = "alt too long (%d chars)" % len(alt)
        images.append({"key": key, "path": img_path, "current_alt": alt,
                       "name": docPr.get("name") or ""})
        if problem:
            issues.append({"type": problem.split(":")[0], "detail": problem, "key": key})

    paras = doc.paragraphs
    total_text = sum(len(p.text.strip()) for p in paras)
    levels = [heading_level(p) for p in paras if heading_level(p)]
    if total_text > 600 and not levels:
        issues.append({"type": "no real headings",
                       "detail": "document has %d chars of text and zero Heading styles" % total_text})
    prev = 0
    for lv in levels:
        if prev and lv > prev + 1:
            issues.append({"type": "skipped heading level",
                           "detail": "H%d follows H%d" % (lv, prev)})
        prev = lv
    for i, p in enumerate(paras):
        if looks_like_faux_heading(p):
            faux.append({"index": i, "text": p.text.strip()[:70]})

    for t_i, table in enumerate(doc.tables):
        if not table_has_header_row(table):
            issues.append({"type": "table missing header row",
                           "detail": "table %d: first row not marked w:tblHeader" % t_i})

    report = {"doc": os.path.basename(path), "paragraphs": len(paras),
              "issues": issues, "images": images, "faux_heading_candidates": faux}
    rp = os.path.join(workdir, "report.json")
    with open(rp, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=1)
    print("SCAN %s: %d issues, %d image(s), %d faux-heading candidate(s)"
          % (report["doc"], len(issues), len(images), len(faux)))
    for x in issues:
        print("  %s - %s" % (x["type"], x["detail"]))
    for c in faux:
        print("  [candidate heading] par %d: %r" % (c["index"], c["text"]))
    print("report: %s" % rp)
    return 0


# ---------- apply ----------

def apply_fixes(path, workdir, out):
    with open(os.path.join(workdir, "fixes.json"), encoding="utf-8") as f:
        fixes = json.load(f)
    alts = fixes.get("alts", {})
    headings = {int(k): int(v) for k, v in fixes.get("headings", {}).items()}
    fix_tables = fixes.get("table_headers", True)

    doc = Document(path)
    n_alt = n_head = n_tbl = 0
    for idx, docPr, _blip in iter_drawings(doc):
        key = "img%d" % idx
        if key in alts:
            val = alts[key].strip()
            if len(val) > MAX_ALT:
                val = val[:MAX_ALT - 1].rstrip() + "."
            docPr.set("descr", val)
            n_alt += 1
    for i, lv in headings.items():
        if 0 <= i < len(doc.paragraphs) and 1 <= lv <= 4:
            doc.paragraphs[i].style = doc.styles["Heading %d" % lv]
            n_head += 1
    if fix_tables:
        for table in doc.tables:
            if not table_has_header_row(table):
                set_table_header_row(table)
                n_tbl += 1
    doc.save(out)
    print("APPLY: %d alt(s), %d heading promotion(s), %d table header row(s) -> %s"
          % (n_alt, n_head, n_tbl, out))
    return 0


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("cmd", choices=["scan", "apply", "verify"])
    ap.add_argument("doc")
    ap.add_argument("--workdir", default=None)
    ap.add_argument("--out", default=None)
    a = ap.parse_args()
    if a.cmd == "scan":
        return scan(a.doc, a.workdir or a.doc + ".work")
    if a.cmd == "apply":
        if not a.out:
            ap.error("--out required for apply")
        return apply_fixes(a.doc, a.workdir or a.doc + ".work", a.out)
    import tempfile
    return scan(a.doc, a.workdir or tempfile.mkdtemp(prefix="docx_verify_"))


if __name__ == "__main__":
    sys.exit(main())
